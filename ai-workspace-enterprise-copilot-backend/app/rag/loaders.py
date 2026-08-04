from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".md",
    ".txt",
}


class DocumentExtractionError(Exception):
    pass


def extract_pdf_text(file_path: Path) -> str:
    reader = PdfReader(str(file_path))

    pages: list[str] = []

    for page_number, page in enumerate(
        reader.pages,
        start=1,
    ):
        page_text = page.extract_text() or ""
        page_text = page_text.strip()

        if page_text:
            pages.append(
                f"[Page {page_number}]\n{page_text}"
            )

    return "\n\n".join(pages)


def extract_docx_text(file_path: Path) -> str:
    document = DocxDocument(str(file_path))

    sections: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            sections.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            ]

            if cells:
                sections.append(" | ".join(cells))

    return "\n\n".join(sections)


def extract_plain_text(file_path: Path) -> str:
    try:
        return file_path.read_text(
            encoding="utf-8"
        ).strip()

    except UnicodeDecodeError:
        return file_path.read_text(
            encoding="latin-1"
        ).strip()


def extract_text(file_path: Path) -> str:
    extension = file_path.suffix.lower()

    try:
        if extension == ".pdf":
            text = extract_pdf_text(file_path)

        elif extension == ".docx":
            text = extract_docx_text(file_path)

        elif extension in {".md", ".txt"}:
            text = extract_plain_text(file_path)

        else:
            raise DocumentExtractionError(
                f"Unsupported file type: {extension}"
            )

    except DocumentExtractionError:
        raise

    except Exception as error:
        raise DocumentExtractionError(
            f"Could not extract text: {error}"
        ) from error

    if not text.strip():
        raise DocumentExtractionError(
            "No readable text was found in the document."
        )

    return text.strip()